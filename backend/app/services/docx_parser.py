"""
DOCX Document Processing Service.
Uses python-docx to extract raw text content from DOCX resumes.
"""

import io
import docx

def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extracts text paragraphs and tables from DOCX file bytes.
    """
    extracted_text = ""
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        for para in doc.paragraphs:
            if para.text:
                extracted_text += para.text + "\n"
                
        # Extract content from tables if present
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    extracted_text += " | ".join(row_text) + "\n"
    except Exception as e:
        print(f"[DOCX Parser Error] {e}")

    return extracted_text.strip()
